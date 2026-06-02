from __future__ import annotations

import json
from pathlib import Path
from typing import Dict, Iterable, Tuple

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def _ensure_reports_dir(reports_dir: str) -> Path:
    output = Path(reports_dir)
    output.mkdir(parents=True, exist_ok=True)
    return output


def _iter_detail_items(details: object) -> Iterable[Tuple[str, list[str]]]:
    if isinstance(details, dict):
        for key, items in details.items():
            title = key.replace("_", " ").title()
            if isinstance(items, dict):
                yield title, [json.dumps(items, indent=2)]
            elif isinstance(items, list):
                yield title, [str(item) for item in items]
            else:
                yield title, [str(items)]
        return
    if isinstance(details, list):
        yield "Details", [str(item) for item in details]


def _summary_rows(result: Dict) -> list[tuple[str, str]]:
    summary = result.get("summary") or {}
    return [
        ("Analysis ID", str(result.get("analysis_id", ""))),
        ("Club", str(summary.get("club", result.get("club", "Unknown")))),
        ("Swing Score", str(summary.get("swing_score", result.get("swing_score", 0)))),
        ("Next Focus", str(summary.get("next_focus", result.get("next_focus", "")))),
    ]


def generate_pdf_report(result: Dict, reports_dir: str) -> str:
    """Generate a local PDF report from analysis results."""
    output_dir = _ensure_reports_dir(reports_dir)
    report_name = f"swing_report_{result['analysis_id']}.pdf"
    report_path = output_dir / report_name

    pdf = canvas.Canvas(str(report_path), pagesize=letter)
    width, height = letter
    y = height - 50

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "SwingSight AI Swing Report")
    y -= 25

    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(50, y, "Summary")
    y -= 15
    pdf.setFont("Helvetica", 10)
    for label, value in _summary_rows(result):
        pdf.drawString(50, y, f"- {label}: {value}")
        y -= 14
        if y < 70:
            pdf.showPage()
            y = height - 50

    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(50, y, "Strengths")
    y -= 15
    pdf.setFont("Helvetica", 10)
    for item in result.get("strengths", []):
        pdf.drawString(50, y, f"- {item}")
        y -= 14
        if y < 70:
            pdf.showPage()
            y = height - 50

    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(50, y, "Improvements")
    y -= 15
    pdf.setFont("Helvetica", 10)
    for item in result.get("improvements", []):
        pdf.drawString(50, y, f"- {item}")
        y -= 14
        if y < 70:
            pdf.showPage()
            y = height - 50

    pdf.setFont("Helvetica-Bold", 11)
    pdf.drawString(50, y, "Advanced Details")
    y -= 15
    pdf.setFont("Helvetica", 10)
    for title, items in _iter_detail_items(result.get("advanced_metrics")):
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(50, y, title)
        y -= 14
        pdf.setFont("Helvetica", 9)
        for item in items:
            text = item if len(item) < 120 else item[:117] + "..."
            pdf.drawString(60, y, f"- {text}")
            y -= 12
            if y < 70:
                pdf.showPage()
                y = height - 50

    pdf.save()
    return str(report_path)


def generate_word_report(result: Dict, reports_dir: str) -> str:
    """Generate a local Word report from analysis results."""
    output_dir = _ensure_reports_dir(reports_dir)
    report_name = f"swing_report_{result['analysis_id']}.docx"
    report_path = output_dir / report_name

    document = Document()
    document.add_heading("SwingSight AI Swing Report", level=1)

    document.add_heading("Summary", level=2)
    for label, value in _summary_rows(result):
        document.add_paragraph(f"{label}: {value}")

    document.add_heading("Strengths", level=2)
    for item in result.get("strengths", []):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Improvements", level=2)
    for item in result.get("improvements", []):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Advanced Details", level=2)
    advanced = result.get("advanced_metrics", {})
    if isinstance(advanced, dict):
        for title, items in _iter_detail_items(advanced):
            document.add_heading(title, level=3)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph("No advanced details available.")

    document.save(str(report_path))
    return str(report_path)